"""
DuckyCoder v6 Universal Input Processor
Handles ingestion and processing of multiple file formats and input types.
"""

import asyncio
import chardet
import mimetypes
import os
import re
import requests
import subprocess
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import logging
import hashlib
import zipfile
import tarfile
from urllib.parse import urlparse
from dataclasses import dataclass
from datetime import datetime

# Format-specific imports
import json
import yaml
import csv
import xml.etree.ElementTree as ET
import base64

# Optional imports for advanced file types
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import nbformat
    HAS_NBFORMAT = True
except ImportError:
    HAS_NBFORMAT = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


@dataclass
class InputMetadata:
    """Metadata for processed input."""
    file_path: str
    file_type: str
    encoding: str
    size: int
    language: Optional[str]
    framework: Optional[str]
    ui_detected: bool
    created_at: datetime
    checksum: str


@dataclass
class ProcessedInput:
    """Processed input data with metadata."""
    content: str
    metadata: InputMetadata
    raw_data: Optional[bytes]
    structured_data: Optional[Dict[str, Any]]
    dependencies: List[str]
    imports: List[str]
    ui_components: List[str]
    error: Optional[str]


class UniversalInputProcessor:
    """Universal input processor supporting 25+ file formats."""
    
    SUPPORTED_CODE_EXTENSIONS = {
        '.py': 'python', '.js': 'javascript', '.ts': 'typescript',
        '.jsx': 'javascript', '.tsx': 'typescript', '.java': 'java',
        '.cs': 'csharp', '.cpp': 'cpp', '.c': 'c', '.rs': 'rust',
        '.go': 'go', '.rb': 'ruby', '.php': 'php', '.sh': 'bash',
        '.ps1': 'powershell', '.kt': 'kotlin', '.swift': 'swift',
        '.dart': 'dart', '.scala': 'scala', '.clj': 'clojure',
        '.hs': 'haskell', '.elm': 'elm', '.ml': 'ocaml', '.fs': 'fsharp',
        '.r': 'r', '.m': 'matlab'
    }
    
    SUPPORTED_MARKUP_EXTENSIONS = {
        '.html': 'html', '.htm': 'html', '.xml': 'xml',
        '.css': 'css', '.scss': 'scss', '.less': 'less',
        '.vue': 'vue', '.svelte': 'svelte'
    }
    
    SUPPORTED_DATA_EXTENSIONS = {
        '.json': 'json', '.yaml': 'yaml', '.yml': 'yaml',
        '.csv': 'csv', '.tsv': 'tsv', '.sql': 'sql'
    }
    
    SUPPORTED_DOCUMENT_EXTENSIONS = {
        '.md': 'markdown', '.txt': 'text', '.pdf': 'pdf',
        '.docx': 'docx', '.ipynb': 'jupyter'
    }
    
    UI_FRAMEWORKS = {
        'react': ['react', 'jsx', 'tsx', 'useState', 'useEffect', 'component'],
        'vue': ['vue', 'template', 'script', 'style', 'v-'],
        'angular': ['angular', '@component', 'ngfor', 'ngif', 'service'],
        'flutter': ['flutter', 'widget', 'stateful', 'stateless', 'scaffold'],
        'tkinter': ['tkinter', 'tk', 'widget', 'mainloop', 'pack', 'grid'],
        'pyqt': ['pyqt', 'qwidget', 'qapplication', 'qmainwindow'],
        'kivy': ['kivy', 'app', 'widget', 'layout', 'canvas'],
        'egui': ['egui', 'ui', 'ctx', 'central_panel', 'side_panel'],
        'dioxus': ['dioxus', 'component', 'rsx!', 'use_state'],
        'tui-rs': ['tui', 'terminal', 'backend', 'frame', 'block']
    }
    
    def __init__(self, config: Dict[str, Any]):
        """Initialize the input processor."""
        self.config = config
        self.input_config = config.get('input_config', {})
        self.max_file_size = self._parse_size(self.input_config.get('max_file_size', '20MB'))
        self.logger = logging.getLogger(__name__)
        
        # Initialize supported formats
        self.all_extensions = {
            **self.SUPPORTED_CODE_EXTENSIONS,
            **self.SUPPORTED_MARKUP_EXTENSIONS,
            **self.SUPPORTED_DATA_EXTENSIONS,
            **self.SUPPORTED_DOCUMENT_EXTENSIONS
        }
    
    async def process_input(self, input_source: str) -> ProcessedInput:
        """
        Process any input source (file, URL, directory, content string).
        
        Args:
            input_source: File path, URL, directory, or content string
            
        Returns:
            ProcessedInput object with content and metadata
        """
        try:
            # Determine input type
            if self._is_url(input_source):
                return await self._process_url(input_source)
            elif os.path.isfile(input_source):
                return await self._process_file(input_source)
            elif os.path.isdir(input_source):
                return await self._process_directory(input_source)
            else:
                # Treat as content string
                return await self._process_content_string(input_source)
                
        except Exception as e:
            self.logger.error(f"Failed to process input {input_source}: {e}")
            return ProcessedInput(
                content="",
                metadata=InputMetadata(
                    file_path=input_source,
                    file_type="unknown",
                    encoding="unknown",
                    size=0,
                    language=None,
                    framework=None,
                    ui_detected=False,
                    created_at=datetime.now(),
                    checksum=""
                ),
                raw_data=None,
                structured_data=None,
                dependencies=[],
                imports=[],
                ui_components=[],
                error=str(e)
            )
    
    async def _process_file(self, file_path: str) -> ProcessedInput:
        """Process a single file."""
        file_path = Path(file_path)
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File size {file_size} exceeds maximum {self.max_file_size}")
        
        # Read file content
        raw_data = file_path.read_bytes()
        checksum = hashlib.sha256(raw_data).hexdigest()
        
        # Detect encoding
        encoding = self._detect_encoding(raw_data)
        
        # Determine file type and language
        file_extension = file_path.suffix.lower()
        file_type = self._determine_file_type(file_path, raw_data)
        language = self._detect_language(file_path, raw_data)
        
        # Process content based on file type
        content, structured_data = await self._process_by_type(
            raw_data, file_type, encoding, file_path
        )
        
        # Extract dependencies and imports
        dependencies = self._extract_dependencies(content, language)
        imports = self._extract_imports(content, language)
        
        # Detect UI components and framework
        ui_components = self._detect_ui_components(content)
        framework = self._detect_framework(content)
        ui_detected = bool(ui_components or framework)
        
        # Create metadata
        metadata = InputMetadata(
            file_path=str(file_path),
            file_type=file_type,
            encoding=encoding,
            size=file_size,
            language=language,
            framework=framework,
            ui_detected=ui_detected,
            created_at=datetime.now(),
            checksum=checksum
        )
        
        return ProcessedInput(
            content=content,
            metadata=metadata,
            raw_data=raw_data,
            structured_data=structured_data,
            dependencies=dependencies,
            imports=imports,
            ui_components=ui_components,
            error=None
        )
    
    async def _process_directory(self, dir_path: str) -> ProcessedInput:
        """Process a directory recursively."""
        dir_path = Path(dir_path)
        all_files = []
        combined_content = []
        combined_dependencies = []
        combined_imports = []
        combined_ui_components = []
        
        # Recursively find all supported files
        if self.input_config.get('recursive_directory_scan', True):
            for file_path in dir_path.rglob('*'):
                if file_path.is_file() and self._is_supported_file(file_path):
                    all_files.append(file_path)
        else:
            for file_path in dir_path.iterdir():
                if file_path.is_file() and self._is_supported_file(file_path):
                    all_files.append(file_path)
        
        # Process each file
        for file_path in all_files:
            try:
                processed_file = await self._process_file(str(file_path))
                if not processed_file.error:
                    combined_content.append(f"# File: {file_path.relative_to(dir_path)}\n")
                    combined_content.append(processed_file.content)
                    combined_content.append("\n\n")
                    
                    combined_dependencies.extend(processed_file.dependencies)
                    combined_imports.extend(processed_file.imports)
                    combined_ui_components.extend(processed_file.ui_components)
                    
            except Exception as e:
                self.logger.warning(f"Failed to process file {file_path}: {e}")
        
        # Combine all content
        final_content = "".join(combined_content)
        
        # Create metadata for directory
        metadata = InputMetadata(
            file_path=str(dir_path),
            file_type="directory",
            encoding="utf-8",
            size=len(final_content.encode('utf-8')),
            language=self._detect_primary_language(all_files),
            framework=self._detect_primary_framework(combined_ui_components),
            ui_detected=bool(combined_ui_components),
            created_at=datetime.now(),
            checksum=hashlib.sha256(final_content.encode('utf-8')).hexdigest()
        )
        
        return ProcessedInput(
            content=final_content,
            metadata=metadata,
            raw_data=final_content.encode('utf-8'),
            structured_data={"files_processed": len(all_files)},
            dependencies=list(set(combined_dependencies)),
            imports=list(set(combined_imports)),
            ui_components=list(set(combined_ui_components)),
            error=None
        )
    
    async def _process_url(self, url: str) -> ProcessedInput:
        """Process content from URL."""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Determine content type
            content_type = response.headers.get('content-type', '').lower()
            
            if 'application/json' in content_type:
                content = json.dumps(response.json(), indent=2)
                file_type = 'json'
            elif 'text/' in content_type or 'application/javascript' in content_type:
                content = response.text
                file_type = self._determine_file_type_from_content(content, url)
            else:
                # Binary content
                content = base64.b64encode(response.content).decode('utf-8')
                file_type = 'binary'
            
            # Create metadata
            metadata = InputMetadata(
                file_path=url,
                file_type=file_type,
                encoding=response.encoding or 'utf-8',
                size=len(response.content),
                language=self._detect_language_from_content(content),
                framework=self._detect_framework(content),
                ui_detected=self._detect_ui_components(content) != [],
                created_at=datetime.now(),
                checksum=hashlib.sha256(response.content).hexdigest()
            )
            
            return ProcessedInput(
                content=content,
                metadata=metadata,
                raw_data=response.content,
                structured_data=None,
                dependencies=self._extract_dependencies(content, metadata.language),
                imports=self._extract_imports(content, metadata.language),
                ui_components=self._detect_ui_components(content),
                error=None
            )
            
        except Exception as e:
            raise ValueError(f"Failed to process URL {url}: {e}")
    
    async def _process_content_string(self, content: str) -> ProcessedInput:
        """Process raw content string."""
        # Try to detect language from content
        language = self._detect_language_from_content(content)
        framework = self._detect_framework(content)
        ui_components = self._detect_ui_components(content)
        
        # Create metadata
        metadata = InputMetadata(
            file_path="<string>",
            file_type="text",
            encoding="utf-8",
            size=len(content.encode('utf-8')),
            language=language,
            framework=framework,
            ui_detected=bool(ui_components),
            created_at=datetime.now(),
            checksum=hashlib.sha256(content.encode('utf-8')).hexdigest()
        )
        
        return ProcessedInput(
            content=content,
            metadata=metadata,
            raw_data=content.encode('utf-8'),
            structured_data=None,
            dependencies=self._extract_dependencies(content, language),
            imports=self._extract_imports(content, language),
            ui_components=ui_components,
            error=None
        )
    
    async def _process_by_type(self, raw_data: bytes, file_type: str, 
                             encoding: str, file_path: Path) -> tuple[str, Optional[Dict[str, Any]]]:
        """Process content based on file type."""
        try:
            if file_type in ['python', 'javascript', 'typescript', 'java', 'cpp', 'c', 'rust', 'go']:
                # Code files - decode as text
                content = raw_data.decode(encoding)
                return content, None
                
            elif file_type == 'json':
                content = raw_data.decode(encoding)
                try:
                    structured_data = json.loads(content)
                    return content, structured_data
                except json.JSONDecodeError:
                    return content, None
                    
            elif file_type in ['yaml', 'yml']:
                content = raw_data.decode(encoding)
                try:
                    structured_data = yaml.safe_load(content)
                    return content, structured_data
                except yaml.YAMLError:
                    return content, None
                    
            elif file_type == 'csv':
                content = raw_data.decode(encoding)
                if HAS_PANDAS:
                    try:
                        df = pd.read_csv(file_path)
                        structured_data = df.to_dict('records')
                        return content, structured_data
                    except Exception:
                        return content, None
                else:
                    return content, None
                    
            elif file_type == 'xml':
                content = raw_data.decode(encoding)
                try:
                    root = ET.fromstring(content)
                    structured_data = self._xml_to_dict(root)
                    return content, structured_data
                except ET.ParseError:
                    return content, None
                    
            elif file_type == 'pdf' and HAS_PDF:
                return await self._process_pdf(raw_data)
                
            elif file_type == 'docx' and HAS_DOCX:
                return await self._process_docx(raw_data)
                
            elif file_type == 'jupyter' and HAS_NBFORMAT:
                return await self._process_jupyter(raw_data, encoding)
                
            else:
                # Default: try to decode as text
                content = raw_data.decode(encoding, errors='ignore')
                return content, None
                
        except Exception as e:
            self.logger.warning(f"Failed to process {file_type}: {e}")
            # Fallback to raw text
            content = raw_data.decode(encoding, errors='ignore')
            return content, None
    
    async def _process_pdf(self, raw_data: bytes) -> tuple[str, Optional[Dict[str, Any]]]:
        """Process PDF file."""
        import io
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(raw_data))
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                text_content.append(f"# Page {page_num + 1}\n{text}\n")
            
            content = "\n".join(text_content)
            structured_data = {
                "page_count": len(pdf_reader.pages),
                "metadata": pdf_reader.metadata
            }
            
            return content, structured_data
            
        except Exception as e:
            raise ValueError(f"Failed to process PDF: {e}")
    
    async def _process_docx(self, raw_data: bytes) -> tuple[str, Optional[Dict[str, Any]]]:
        """Process DOCX file."""
        import io
        
        try:
            doc = Document(io.BytesIO(raw_data))
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            content = "\n".join(text_content)
            structured_data = {
                "paragraph_count": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return content, structured_data
            
        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {e}")
    
    async def _process_jupyter(self, raw_data: bytes, encoding: str) -> tuple[str, Optional[Dict[str, Any]]]:
        """Process Jupyter notebook."""
        try:
            content = raw_data.decode(encoding)
            notebook = nbformat.reads(content, as_version=4)
            
            text_content = []
            code_cells = []
            markdown_cells = []
            
            for cell in notebook.cells:
                if cell.cell_type == 'code':
                    text_content.append(f"```{notebook.metadata.get('kernelspec', {}).get('language', 'python')}")
                    text_content.append(cell.source)
                    text_content.append("```\n")
                    code_cells.append(cell.source)
                elif cell.cell_type == 'markdown':
                    text_content.append(cell.source)
                    text_content.append("\n")
                    markdown_cells.append(cell.source)
            
            processed_content = "\n".join(text_content)
            structured_data = {
                "cell_count": len(notebook.cells),
                "code_cells": len(code_cells),
                "markdown_cells": len(markdown_cells),
                "language": notebook.metadata.get('kernelspec', {}).get('language', 'python')
            }
            
            return processed_content, structured_data
            
        except Exception as e:
            raise ValueError(f"Failed to process Jupyter notebook: {e}")
    
    def _detect_encoding(self, raw_data: bytes) -> str:
        """Detect file encoding."""
        if self.input_config.get('auto_convert_encoding', True):
            try:
                result = chardet.detect(raw_data)
                return result['encoding'] or 'utf-8'
            except Exception:
                return 'utf-8'
        else:
            return 'utf-8'
    
    def _determine_file_type(self, file_path: Path, raw_data: bytes) -> str:
        """Determine file type from extension and content."""
        extension = file_path.suffix.lower()
        
        # Check known extensions
        if extension in self.all_extensions:
            return self.all_extensions[extension]
        
        # AI-powered filetype detection
        if self.input_config.get('ai_filetype_detection', True):
            return self._ai_detect_filetype(raw_data, str(file_path))
        
        # Fallback to mimetype
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if mime_type:
            if mime_type.startswith('text/'):
                return 'text'
            elif mime_type.startswith('application/'):
                return mime_type.split('/')[-1]
        
        return 'unknown'
    
    def _determine_file_type_from_content(self, content: str, url: str) -> str:
        """Determine file type from content and URL."""
        url_path = urlparse(url).path
        extension = Path(url_path).suffix.lower()
        
        if extension in self.all_extensions:
            return self.all_extensions[extension]
        
        # Analyze content patterns
        if content.strip().startswith('{') and content.strip().endswith('}'):
            return 'json'
        elif '<!DOCTYPE html' in content or '<html' in content:
            return 'html'
        elif content.strip().startswith('<?xml'):
            return 'xml'
        
        return 'text'
    
    def _detect_language(self, file_path: Path, raw_data: bytes) -> Optional[str]:
        """Detect programming language."""
        extension = file_path.suffix.lower()
        
        if extension in self.SUPPORTED_CODE_EXTENSIONS:
            return self.SUPPORTED_CODE_EXTENSIONS[extension]
        
        # Try content-based detection
        try:
            content = raw_data.decode('utf-8', errors='ignore')[:1000]  # First 1KB
            return self._detect_language_from_content(content)
        except Exception:
            return None
    
    def _detect_language_from_content(self, content: str) -> Optional[str]:
        """Detect language from content patterns."""
        content_lower = content.lower()
        
        # Language patterns
        patterns = {
            'python': ['def ', 'import ', 'from ', 'class ', '__init__', 'if __name__'],
            'javascript': ['function ', 'const ', 'let ', 'var ', '=>', 'console.log'],
            'typescript': ['interface ', 'type ', ': string', ': number', 'export'],
            'java': ['public class', 'private ', 'public static void main'],
            'rust': ['fn ', 'let mut', 'impl ', 'struct ', 'enum '],
            'go': ['func ', 'package ', 'import ', 'var ', 'type '],
            'cpp': ['#include', 'std::', 'namespace ', 'class '],
            'bash': ['#!/bin/bash', 'echo ', 'if [', 'then', 'fi']
        }
        
        scores = {}
        for lang, lang_patterns in patterns.items():
            score = sum(1 for pattern in lang_patterns if pattern in content_lower)
            if score > 0:
                scores[lang] = score
        
        if scores:
            return max(scores, key=scores.get)
        
        return None
    
    def _detect_framework(self, content: str) -> Optional[str]:
        """Detect UI framework from content."""
        content_lower = content.lower()
        
        for framework, patterns in self.UI_FRAMEWORKS.items():
            if any(pattern in content_lower for pattern in patterns):
                return framework
        
        return None
    
    def _detect_ui_components(self, content: str) -> List[str]:
        """Detect UI components in content."""
        components = []
        content_lower = content.lower()
        
        # Common UI component patterns
        ui_patterns = [
            'button', 'input', 'form', 'dialog', 'modal', 'menu',
            'window', 'panel', 'widget', 'component', 'layout',
            'grid', 'list', 'table', 'canvas', 'frame'
        ]
        
        for pattern in ui_patterns:
            if pattern in content_lower:
                components.append(pattern)
        
        return components
    
    def _extract_dependencies(self, content: str, language: Optional[str]) -> List[str]:
        """Extract dependencies from content."""
        dependencies = []
        
        if language == 'python':
            # Python imports
            import_patterns = [
                r'import\s+([^\s]+)',
                r'from\s+([^\s]+)\s+import'
            ]
        elif language in ['javascript', 'typescript']:
            # JS/TS imports
            import_patterns = [
                r'import.*from\s+[\'"]([^\'"]+)[\'"]',
                r'require\([\'"]([^\'"]+)[\'"]\)'
            ]
        elif language == 'java':
            # Java imports
            import_patterns = [r'import\s+([^\s;]+)']
        elif language == 'rust':
            # Rust uses
            import_patterns = [r'use\s+([^\s;]+)']
        elif language == 'go':
            # Go imports
            import_patterns = [r'import\s+"([^"]+)"']
        else:
            return dependencies
        
        for pattern in import_patterns:
            matches = re.findall(pattern, content)
            dependencies.extend(matches)
        
        return dependencies
    
    def _extract_imports(self, content: str, language: Optional[str]) -> List[str]:
        """Extract import statements from content."""
        imports = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if language == 'python' and (line.startswith('import ') or line.startswith('from ')):
                imports.append(line)
            elif language in ['javascript', 'typescript'] and 'import' in line:
                imports.append(line)
            elif language == 'java' and line.startswith('import '):
                imports.append(line)
            elif language == 'rust' and line.startswith('use '):
                imports.append(line)
            elif language == 'go' and 'import' in line:
                imports.append(line)
        
        return imports
    
    def _ai_detect_filetype(self, raw_data: bytes, file_path: str) -> str:
        """AI-powered filetype detection."""
        # Simplified heuristic-based detection
        try:
            content = raw_data.decode('utf-8', errors='ignore')[:500]
            
            # Check for common patterns
            if content.strip().startswith('{') and content.strip().endswith('}'):
                return 'json'
            elif content.strip().startswith('---') or 'key:' in content:
                return 'yaml'
            elif '#!/bin/bash' in content or '#!/bin/sh' in content:
                return 'bash'
            elif '#!/usr/bin/env python' in content or 'def ' in content:
                return 'python'
            elif 'function ' in content or 'const ' in content:
                return 'javascript'
            
        except Exception:
            pass
        
        return 'text'
    
    def _xml_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary."""
        result = {}
        
        # Add attributes
        if element.attrib:
            result['@attributes'] = element.attrib
        
        # Add text content
        if element.text and element.text.strip():
            if len(element) == 0:
                return element.text.strip()
            result['text'] = element.text.strip()
        
        # Add children
        for child in element:
            child_data = self._xml_to_dict(child)
            if child.tag in result:
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_data)
            else:
                result[child.tag] = child_data
        
        return result
    
    def _detect_primary_language(self, files: List[Path]) -> Optional[str]:
        """Detect primary language from file list."""
        language_counts = {}
        
        for file_path in files:
            extension = file_path.suffix.lower()
            if extension in self.SUPPORTED_CODE_EXTENSIONS:
                lang = self.SUPPORTED_CODE_EXTENSIONS[extension]
                language_counts[lang] = language_counts.get(lang, 0) + 1
        
        if language_counts:
            return max(language_counts, key=language_counts.get)
        
        return None
    
    def _detect_primary_framework(self, ui_components: List[str]) -> Optional[str]:
        """Detect primary framework from UI components."""
        if not ui_components:
            return None
        
        # Simple heuristic - return most common framework indicator
        framework_counts = {}
        for component in ui_components:
            for framework, patterns in self.UI_FRAMEWORKS.items():
                if component in patterns:
                    framework_counts[framework] = framework_counts.get(framework, 0) + 1
        
        if framework_counts:
            return max(framework_counts, key=framework_counts.get)
        
        return None
    
    def _is_url(self, input_source: str) -> bool:
        """Check if input is a URL."""
        return input_source.startswith(('http://', 'https://'))
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file is supported."""
        extension = file_path.suffix.lower()
        return extension in self.all_extensions
    
    def _parse_size(self, size_str: str) -> int:
        """Parse size string to bytes."""
        size_str = size_str.upper()
        if size_str.endswith('KB'):
            return int(size_str[:-2]) * 1024
        elif size_str.endswith('MB'):
            return int(size_str[:-2]) * 1024 * 1024
        elif size_str.endswith('GB'):
            return int(size_str[:-2]) * 1024 * 1024 * 1024
        else:
            return int(size_str)